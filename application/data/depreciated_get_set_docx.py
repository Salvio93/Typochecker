
import docx
from .manage_document_rule_db import *
from .manage_document_db import *
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from copy import deepcopy




def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        """
        for run in para.runs:
            if run.bold:
                print (run.text +' bold')
            if run.italic:
                print(run.text + ' italic')
            if run.font.size != Pt(11):
                print(run.text + ' font size not 11')
            if run.font.name != "Sofia" or run.font.name != "Multi":
                print(run.text + ' font policy not correct')
        """
        fullText.append(para.text)
    return '\n'.join(fullText)

def setText(text,filename):
    doc = docx.Document(filename)
    list_of_para =text.split('\n')
    
    for idx,para in enumerate(doc.paragraphs):
        if (idx < len(list_of_para)):
            para.text=list_of_para[idx]

    doc.save(filename)
    #set_all_hightlight_of_doc(filename)


def insert_run_at_position(par, pos, txt=''):
    """Insert a new run with text {txt} into paragraph {par}
    at given position {pos}.

    Returns the newly created run.
    """
    p = par._p
    new_run = par.add_run(txt)
    p.insert(pos + 1, new_run._r)

    return new_run


def insert_run_before(par, run, txt=''):
    """Insert a new run with text {txt} into paragraph before given {run}.

    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addprevious(run_2._r)

    return run_2


def insert_run_after(par, run, txt=''):
    """Insert a new run with text {txt} into paragraph after given {run}.

    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addnext(run_2._r)

    return run_2

def copy_run_format(run_src, run_dst):
    """Copy formatting from {run_src} to {run_dst}.
    """
    rPr_target = run_dst._r.get_or_add_rPr()
    rPr_target.addnext(deepcopy(run_src._r.get_or_add_rPr()))
    run_dst._r.remove(rPr_target)


def split_run_by(par, run, split_by):
    """Split text in {run} from paragraph {par} by positions
    provided by {split_by}, while retaining original {run}
    formatting.

    Returns list of split runs starting with original {run}.
    """
    txt = run.text
    txt_len = len(txt)
    if not all(isinstance(i, int) for i in split_by):
        raise ValueError("Split positions must be integer numbers")
    split_list = [i if i >= 0 else txt_len + i for i in split_by]
    if not all(split_list[j] <= split_list[j + 1]
               for j in range(len(split_list) - 1)):
        raise ValueError("Split positions must be sorted to make sense")
    if split_list[0] < 0:
        raise ValueError("A split position cannot be less than -<text length>")
    split_list.insert(0, 0)
    split_list.append(None)
    split_txts = [txt[split_list[i]:split_list[i + 1]]
                  for i in range(len(split_list) - 1)]
    run.text = split_txts[0]
    split_txts.pop(0)
    new_runs = [run]
    for next_txt in split_txts:
        new_runs.append(insert_run_after(par, new_runs[-1], next_txt))
        copy_run_format(run, new_runs[-1])

    return new_runs



def set_high(filename,doc_rule_id):
        
    doc = docx.Document(filename)

    
    indexage = 0
    ite=0
    apply_rules= retrieve_docu_rule_list(doc_rule_id)[0]
    error_idx = apply_rules[6]
    error_len = apply_rules[7]
    
    for para in doc.paragraphs:
        
        for run in para.runs:
            run_len = len(run.text)
            
            #CA FAIT FONCTIONNER LE CODE, LE FAIT DE SPLIT_RUN CREE DES RUN VIDE ET CA CREAIT DES PROBLEMES D'INDEXAGE POUR LES PROCHAINS SPLIT
            if run_len ==0:
                ite-=1
            
            else:
                if indexage <= error_idx and indexage+ run_len >= error_idx:
                    #print(run.text[error_idx-indexage-ite:error_idx+error_len-indexage-ite])
                    #print(error_idx-indexage-ite,error_idx+error_len-indexage-ite)               
                
                    new_runs = split_run_by(para, run, (error_idx-indexage-ite, error_idx+error_len-indexage-ite))
                    new_runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                indexage+= run_len
        ite+=1    
    doc.save(filename)

def set_all_hightlight_of_doc(file_id):
    doc = retrieve_doc_list(file_id)[0]
    for rule in retrieve_docu_rule_list_by_document_id(doc[0]):
        if rule[-1]==1:
            set_high(doc[-1],rule[0])
            



"""
from spellchecker import SpellChecker

spell = SpellChecker(language='fr')

#add word
spell.word_frequency.load_words(['Isatis'])
# find those words that may be misspelled
misspelled = spell.unknown(['salu'])

for word in misspelled:
    # Get the one `most likely` answer
    print(spell.correction(word))

    # Get a list of `likely` options
    print(spell.candidates(word))
"""

